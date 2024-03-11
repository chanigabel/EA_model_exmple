import win32com.client
import xml.etree.ElementTree as ET
import os
from docx import Document

# Define the path to the SQL query file
sql_file_path = os.path.join(os.path.dirname(__file__), 'query.sql')

# Function to read SQL query from file
def read_sql_query(file_path):
    with open(file_path, 'r') as file:
        sql_query = file.read()
    return sql_query

# Function to generate documents for use case scenarios
def document_use_case_scenarios(repository):
    docgen = repository.CreateDocumentGenerator()
    docgen.NewDocument("UC_ScenarioSteps")
    scenarios = get_scenarios_for_use_case(repository)
    rtf_content = ""
    for scenario in scenarios:
        scenario_xml_string = get_scenario_xml(scenario)
        docgen.DocumentCustomData(scenario_xml_string, 1, "UC_Scenario")
        scenario_steps_xml_string = get_scenario_steps_xml_string(scenario.Name, repository)
        docgen.DocumentCustomData(scenario_steps_xml_string, 1, "UC_ScenarioSteps")
    rtf_content = docgen.GetDocumentAsRTF()
    return rtf_content

# Function to get scenarios for a use case
def get_scenarios_for_use_case(repository):
    scenarios = []
    sql_query = read_sql_query(sql_file_path)
    result_set = repository.SQLQuery(sql_query)
    
    if isinstance(result_set, str):
        root = ET.fromstring(result_set)
    else:
        # If result_set is not a string, convert it to string and then parse
        result_set_str = result_set.decode('utf-8')  # Assuming UTF-8 encoding
        root = ET.fromstring(result_set_str)

    rows = root.findall(".//Row")

    main_scenario_xml_dom = None
    for scenario_row in rows:
        # Extracting data from XML elements in the row
        scenario_name = scenario_row.find("Name").text
        scenario_type = scenario_row.find("CLASSTYPE").text
        xml_content = ET.tostring(scenario_row, encoding='unicode')
        notes = scenario_row.find("Massege").text  # Assuming 'Massege' holds the notes
        guid = scenario_row.find("Guid_Classifier").text
        
        scenario = UsecaseScenario(scenario_name, scenario_type, xml_content, notes, guid)
        scenarios.append(scenario)
        if main_scenario_xml_dom is None:
            main_scenario_xml_dom = ET.fromstring(ET.tostring(scenario_row[2], encoding='unicode'))
        else:
            scenario.resolve_entry_and_join(main_scenario_xml_dom)
    
    return scenarios

# Class representing a use case scenario
class UsecaseScenario:
    def __init__(self, name, scenario_type, xml_content, notes, guid):
        self.Name = name
        self.Notes = notes
        self.ScenarioType = scenario_type
        self.GUID = guid
        self.Entry = ""
        self.Join = ""
        self.XMLContent = xml_content

    def resolve_entry_and_join(self, main_scenario_xml_dom):
        extension_node = main_scenario_xml_dom.find(".//extension[@guid='" + str(self.GUID) + "']")
        if extension_node is not None:
            # Perform further operations only if extension_node is found
            self.Entry = extension_node.get("level")
            join_step_guid = extension_node.get("join")
            if join_step_guid:
                join_step_node = main_scenario_xml_dom.find(".//step[@guid='" + join_step_guid + "']")
                if join_step_node is not None:
                    self.Join = join_step_node.get("level")

# Function to generate XML string for scenario
def get_scenario_xml(scenario):
    xml_root = ET.Element("EADATA")
    dataset = ET.SubElement(xml_root, "Dataset_0")
    data = ET.SubElement(dataset, "Data")
    row = ET.SubElement(data, "Row")

    xml_scenario_type = ET.SubElement(row, "ScenarioType")
    xml_scenario_type.text = scenario.ScenarioType

    xml_name = ET.SubElement(row, "name")
    xml_name.text = scenario.Name

    xml_entry = ET.SubElement(row, "Entry")
    xml_entry.text = scenario.Entry

    xml_join = ET.SubElement(row, "Join")
    xml_join.text = scenario.Join

    xml_notes = ET.SubElement(row, "Notes")
    xml_notes.text = scenario.Notes
    xml_notes.set("formatted", "1")

    return ET.tostring(xml_root, encoding="unicode")

# Function to generate XML string for scenario steps
def get_scenario_steps_xml_string( scenario_name, repository):
    xml_root = ET.Element("EADATA")
    dataset = ET.SubElement(xml_root, "Dataset_0")
    data = ET.SubElement(dataset, "Data")
    add_scenario_contents(data, scenario_name, repository)
    return ET.tostring(xml_root, encoding="unicode")

# Function to add scenario contents
def add_scenario_contents(xml_data, scenario_name, repository):
    # Mock data for testing
    mock_data = [
        (1, 'Step 1', 'Test Scenario'),
        (2, 'Step 2', 'Test Scenario'),
        # Add more rows as needed
    ]

    for row in mock_data:
        step_node = ET.SubElement(xml_data, "Step")
        
        # Add details for each step
        step_number = ET.SubElement(step_node, "StepNumber")
        step_number.text = str(row[0])  # Convert integer to string
        
        step_name = ET.SubElement(step_node, "StepName")
        step_name.text = row[1]
        
        # Add other step details as needed based on your database structure

# Function to insert RTF content into a Word document
def insert_rtf_into_word(rtf_content, output_path):
    # Create a new Word document
    doc = Document()
    
    # Insert the RTF content
    doc.add_paragraph(rtf_content)
    
    # Save the document
    doc.save(output_path)
    print(f"Word document saved at: {output_path}")

# Main function to demonstrate usage
def main():
    ea = win32com.client.Dispatch("EA.App")
    repository = ea.Repository
    repository.OpenFile("C:/Users/ch058/OneDrive/שולחן העבודה/EA_GIT/EA_model_exmple/project_model.eapx")

    # Example: Generate documents for a specific use case object ID
    rtf_content = document_use_case_scenarios(repository)
    
    # Example: Path to save the output Word document
    output_path = "C:/Users/ch058/OneDrive/שולחן העבודה/EA_GIT/output_document.docx"

    # Insert the RTF content into the Word document
    insert_rtf_into_word(rtf_content, output_path)

if __name__ == "__main__":
    main()
